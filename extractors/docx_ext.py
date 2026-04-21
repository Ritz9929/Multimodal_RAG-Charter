"""
DOCX Extractor
===============
Uses python-docx for Word document parsing.

Strategy:
  - Heading-aware sectioning: Heading 1/2 → new "page" break
  - Tables → markdown format (preserving column structure)
  - Images → extracted from media parts, saved to disk
  - Heading hierarchy preserved with markdown markers (# ## ###)

Why heading-based pages?
  Unlike PDFs and PPTXs, DOCX files are a continuous stream.
  Heading breaks provide semantically meaningful section boundaries
  that align with how humans organize documents.
"""

import logging
from pathlib import Path

from docx import Document as DocxDocument
from .base import BaseExtractor, ExtractionResult, ExtractedImage, ExtractedTable

logger = logging.getLogger(__name__)


class DOCXExtractor(BaseExtractor):
    """
    DOCX extraction using python-docx.

    Strategy:
      - Heading-aware sectioning: Heading 1/2 → new "page" break
      - Tables → markdown format (preserving column structure)
      - Images → extracted from media parts, saved to disk
    """

    # Heading styles that trigger a page/section break
    SECTION_BREAK_STYLES = {"Heading 1", "Heading 2"}

    def __init__(self, output_dir: str = "./mock_s3_storage", **kwargs):
        super().__init__(output_dir)

    def _table_to_markdown(self, table) -> str:
        """Convert a python-docx Table object to markdown format."""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)

        if not rows:
            return ""

        # First row as header
        header = "| " + " | ".join(rows[0]) + " |"
        separator = "| " + " | ".join(["---"] * len(rows[0])) + " |"
        body_lines = []
        for row in rows[1:]:
            # Pad row if fewer cells than header
            padded = row + [""] * (len(rows[0]) - len(row))
            body_lines.append("| " + " | ".join(padded[:len(rows[0])]) + " |")

        return "\n".join([header, separator] + body_lines)

    def _extract_images(self, doc: DocxDocument, result: ExtractionResult,
                        doc_output_dir: Path) -> None:
        """Extract all embedded images from the DOCX."""
        img_count = 0
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                try:
                    image_data = rel.target_part.blob
                    if len(image_data) < 1024:
                        continue  # Skip tiny icons

                    # Determine extension from content type
                    content_type = rel.target_part.content_type
                    ext_map = {
                        "image/png": "png",
                        "image/jpeg": "jpeg",
                        "image/jpg": "jpg",
                        "image/gif": "gif",
                        "image/bmp": "bmp",
                        "image/tiff": "tiff",
                    }
                    ext = ext_map.get(content_type, "png")

                    filename = f"docx_img{img_count}.{ext}"
                    filepath = doc_output_dir / filename

                    with open(filepath, "wb") as f:
                        f.write(image_data)

                    result.images.append(ExtractedImage(
                        filename=filename,
                        filepath=str(filepath),
                        page_number=0,  # DOCX doesn't have real page numbers
                        position_index=img_count,
                        source_type="embedded",
                    ))
                    img_count += 1
                    logger.info(f"  Saved DOCX image: {filename}")
                except Exception as e:
                    logger.warning(f"  Skipping DOCX image: {e}")

    def extract(self, file_path: str, source_doc: str = "") -> ExtractionResult:
        """Parse the DOCX and return text + images + tables."""
        doc = DocxDocument(file_path)
        result = ExtractionResult(source_format="docx")

        if not source_doc:
            source_doc = Path(file_path).name
        doc_output_dir = self._make_doc_output_dir(source_doc)

        current_section = 0
        current_text_parts = []
        table_index_on_page = 0

        # Iterate body elements in document order
        for element in doc.element.body:
            tag = element.tag

            if tag.endswith('}p'):
                # ── Paragraph ──
                for para in doc.paragraphs:
                    if para._element is element:
                        style_name = para.style.name if para.style else ""
                        text = para.text.strip()

                        # Check for section break
                        if style_name in self.SECTION_BREAK_STYLES and current_text_parts:
                            result.page_texts[current_section] = "\n".join(current_text_parts)
                            current_text_parts = []
                            current_section += 1
                            table_index_on_page = 0

                        if text:
                            # Add heading markers for context
                            if "Heading 1" in style_name:
                                current_text_parts.append(f"# {text}")
                            elif "Heading 2" in style_name:
                                current_text_parts.append(f"## {text}")
                            elif "Heading 3" in style_name:
                                current_text_parts.append(f"### {text}")
                            else:
                                current_text_parts.append(text)
                        break

            elif tag.endswith('}tbl'):
                # ── Table ──
                for table in doc.tables:
                    if table._element is element:
                        md = self._table_to_markdown(table)
                        if md:
                            current_text_parts.append(md)
                            result.tables.append(ExtractedTable(
                                markdown=md,
                                page_number=current_section,
                                position_index=table_index_on_page,
                                row_count=len(table.rows),
                                col_count=len(table.columns),
                            ))
                            table_index_on_page += 1
                        break

        # Flush remaining text
        if current_text_parts:
            result.page_texts[current_section] = "\n".join(current_text_parts)

        # Extract images
        self._extract_images(doc, result, doc_output_dir)

        result.metadata = {
            "sections": current_section + 1,
            "tables_found": len(result.tables),
            "images_found": len(result.images),
            "total_paragraphs": len(doc.paragraphs),
        }

        logger.info(
            f"DOCX extraction complete — {len(result.page_texts)} sections, "
            f"{len(result.tables)} tables, {len(result.images)} images"
        )
        return result
